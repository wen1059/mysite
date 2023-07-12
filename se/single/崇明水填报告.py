# -*- coding: utf-8 -*-
# date: 2022/7/8
import docx.document
from docx import Document
import os
import xlwings
from xlwings.main import Book, Sheet

os.chdir(r"C:\Users\Administrator\Desktop\崇明水")


def singleline():
    app = xlwings.App(add_book=False, visible=False)
    file: Book = app.books.open('崇明水.22.7.8(1).xlsx')
    sht: Sheet = file.sheets[0]
    for i in range(375, 591):
        # [采样日期	新监测点名称	pH值	高锰酸盐指数	氨氮(NH₃-N)	总磷(TP，以P计)	总氮	样品编号	样品状态	报告编号	报告日期	分析结束日期	pH值设备编号]
        # cyrq, jcdw, ph, mn, cod, nh3, tp, zd, ypbh, ypzt, bgbh, bgrq, fxrq, sbbh
        yield sht.range(f'a{i}:n{i}').value
    app.quit()


def replace():
    doc: docx.document.Document = Document("cod.docx")
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        # print(run.text)
                        if (bgbh_ := '1111A11111111Z') in run.text:
                            run.text = run.text.replace(bgbh_, bgbh)
                        elif (bgrq_ := '0000-00-00') in run.text:
                            run.text = run.text.replace(bgrq_, bgrq)
                        elif (cyrq_ := '1111-11-11') in run.text:
                            run.text = run.text.replace(cyrq_, cyrq)
                        elif (jcrq_ := '2222-22-22') in run.text:
                            run.text = run.text.replace(jcrq_, fxrq)
                        elif (sbbh_ := 'SEMTEC-000') in run.text:
                            run.text = run.text.replace(sbbh_, sbbh)
                        elif (JCDWMC1 := 'CDWMC1') in run.text:
                            run.text = run.text.replace(JCDWMC1, jcdw)
                        elif (YPBH1 := 'PBH1') in run.text:
                            run.text = run.text.replace(YPBH1, ypbh)
                        elif (YPZT1 := 'PZT1') in run.text:
                            run.text = run.text.replace(YPZT1, ypzt)
                        elif (PHZHI1 := 'HZHI1') in run.text:
                            run.text = run.text.replace(PHZHI1, ph)
                        elif (ANDAN1 := 'NDAN1') in run.text:
                            run.text = run.text.replace(ANDAN1, nh3)
                        elif (ZONGLING1 := 'ONGLING1') in run.text:
                            run.text = run.text.replace(ZONGLING1, tp)
                        elif (GMSY1 := 'MSY1') in run.text:
                            run.text = run.text.replace(GMSY1, mn)
                        elif (ZONGDAN1 := 'ONGDAN1') in run.text:
                            run.text = run.text.replace(ZONGDAN1, zd)
                        elif (HXXYL1 := 'XXYL1') in run.text:
                            run.text = run.text.replace(HXXYL1, cod)
    doc.save(f'{bgbh}.docx')


if __name__ == '__main__':
    for cyrq, jcdw, ph, mn, cod, nh3, tp, zd, ypbh, ypzt, bgbh, bgrq, fxrq, sbbh in singleline():
        cyrq = str(cyrq).replace(' 00:00:00', '')
        bgrq = str(bgrq).replace(' 00:00:00', '')
        fxrq = cyrq + '~' + str(fxrq).replace(' 00:00:00', '')
        try:
            cod = str(int(cod))
        except Exception:
            pass
        nh3 = str(nh3)
        # print(cyrq, jcdw, ph, mn, cod, nh3, tp, zd, ypbh, ypzt, bgbh, bgrq, fxrq, sbbh)
        print(ypbh)
        replace()
